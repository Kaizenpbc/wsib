import httpx
import re
from typing import Dict, List
from PyPDF2 import PdfReader
from io import BytesIO
from app.models import Clause, ClauseCategory, Priority
from app.services.zip_extractor import ZipExtractor

class RfpParser:
    """
    Service to parse RFP documents and extract clauses
    """
    
    PRIORITY_KEYWORDS = {
        'must': ['must', 'shall', 'required', 'mandatory'],
        'should': ['should', 'recommended', 'preferred'],
        'may': ['may', 'optional', 'could']
    }
    
    CATEGORY_KEYWORDS = {
        'duration': ['hour', 'minute', 'day', 'week', 'duration', 'time'],
        'content': ['topic', 'subject', 'content', 'cover', 'include', 'teach'],
        'assessment': ['test', 'exam', 'assess', 'evaluat', 'quiz', 'certification'],
        'equipment': ['equipment', 'material', 'tool', 'resource', 'supply']
    }
    
    async def parse_document(self, file_url: str, file_name: str) -> Dict:
        """
        Parse an RFP document and extract clauses
        """
        # Download file
        async with httpx.AsyncClient() as client:
            response = await client.get(file_url)
            file_content = BytesIO(response.content)
        
        # Check if it's a ZIP file
        if file_name.lower().endswith('.zip'):
            return await self._parse_zip(file_content, file_name)
        
        # Determine file type and parse
        if file_name.lower().endswith('.pdf'):
            return await self._parse_pdf(file_content)
        elif file_name.lower().endswith(('.doc', '.docx')):
            return await self._parse_word(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_name}")
    
    async def _parse_pdf(self, file_content: BytesIO) -> Dict:
        """
        Parse PDF document
        """
        reader = PdfReader(file_content)
        total_pages = len(reader.pages)
        
        clauses = []
        
        # Extract text from all pages
        full_text = ""
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            full_text += f"\n[PAGE {page_num}]\n{text}"
        
        # Split into sentences/requirements
        sentences = self._split_into_sentences(full_text)
        
        for sentence in sentences:
            # Only process meaningful requirements
            if len(sentence.strip()) < 20:
                continue
            
            if self._is_requirement(sentence):
                clause = self._create_clause(sentence)
                clauses.append(clause)
        
        return {
            'clauses': clauses,
            'total_pages': total_pages
        }
    
    async def _parse_word(self, file_content: BytesIO) -> Dict:
        """
        Parse Word document
        """
        # Import here to avoid dependency issues
        from docx import Document
        
        doc = Document(file_content)
        clauses = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) < 20:
                continue
                
            if self._is_requirement(text):
                clause = self._create_clause(text)
                clauses.append(clause)
        
        return {
            'clauses': clauses,
            'total_pages': len(doc.paragraphs) // 10  # Rough estimate
        }
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences
        """
        # Simple sentence splitting
        sentences = re.split(r'[.!?]\s+|\n+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _is_requirement(self, text: str) -> bool:
        """
        Determine if text is a requirement
        """
        text_lower = text.lower()
        
        # Check for requirement keywords
        for keywords in self.PRIORITY_KEYWORDS.values():
            for keyword in keywords:
                if keyword in text_lower:
                    return True
        
        # Check for numbered items
        if re.match(r'^\d+[\.\)]\s+', text):
            return True
        
        return False
    
    def _create_clause(self, text: str) -> Clause:
        """
        Create a Clause object from text
        """
        priority = self._determine_priority(text)
        category = self._determine_category(text)
        
        return Clause(
            text=text,
            priority=priority,
            category=category
        )
    
    def _determine_priority(self, text: str) -> Priority:
        """
        Determine the priority of a requirement
        """
        text_lower = text.lower()
        
        for priority, keywords in self.PRIORITY_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return Priority(priority)
        
        return Priority.MAY
    
    def _determine_category(self, text: str) -> ClauseCategory:
        """
        Determine the category of a requirement
        """
        text_lower = text.lower()
        
        scores = {category: 0 for category in self.CATEGORY_KEYWORDS.keys()}
        
        for category, keywords in self.CATEGORY_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    scores[category] += 1
        
        # Return category with highest score
        if max(scores.values()) > 0:
            return ClauseCategory(max(scores, key=scores.get))
        
        return ClauseCategory.OTHER
    
    async def _parse_zip(self, file_content: BytesIO, file_name: str) -> Dict:
        """
        Parse ZIP file containing multiple documents
        """
        import tempfile
        import zipfile
        
        all_clauses = []
        
        try:
            # Save ZIP to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.zip') as temp_file:
                temp_file.write(file_content.read())
                temp_path = temp_file.name
            
            # Extract and process each file
            with zipfile.ZipFile(temp_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                
                for file_name_in_zip in file_list:
                    # Skip directories and hidden files
                    if file_name_in_zip.endswith('/') or file_name_in_zip.startswith('__'):
                        continue
                    
                    file_ext = file_name_in_zip.lower().split('.')[-1]
                    
                    # Process supported file types
                    if file_ext in ['pdf', 'doc', 'docx']:
                        with zip_ref.open(file_name_in_zip) as file_in_zip:
                            file_content_inner = BytesIO(file_in_zip.read())
                            
                            try:
                                if file_ext == 'pdf':
                                    result = await self._parse_pdf(file_content_inner)
                                elif file_ext in ['doc', 'docx']:
                                    result = await self._parse_word(file_content_inner)
                                else:
                                    continue
                                
                                # Add clauses from this file
                                for clause in result['clauses']:
                                    clause.section = f"From: {file_name_in_zip}"
                                    all_clauses.append(clause)
                            
                            except Exception as e:
                                print(f"Error parsing {file_name_in_zip}: {e}")
                                continue
            
            # Clean up temp file
            import os
            os.unlink(temp_path)
            
            return {
                'clauses': all_clauses,
                'total_pages': None,
                'files_processed': len([f for f in file_list if not f.endswith('/')])
            }
        
        except zipfile.BadZipFile:
            raise ValueError("Invalid ZIP file")
        except Exception as e:
            raise ValueError(f"Error processing ZIP file: {str(e)}")

